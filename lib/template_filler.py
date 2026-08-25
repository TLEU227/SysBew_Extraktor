# ============================================================
# template_filler.py
# Erzeugt eine neue Systembewertung (V11) aus einem Daten-Dict - 1.5
#
# Gegenstueck zu den extract_*-Funktionen in sysbew_common.py: dort
# werden Werte AUS einem ausgefuellten Dokument GELESEN, hier werden
# Werte IN das Leer-Template (assets/templates_docx/
# Systembewertung_V11_leer.docx) GESCHRIEBEN.
#
# Wichtiger Unterschied zur Extraktion: die Extraktion muss mit vielen
# unterschiedlichen, bereits ausgefuellten Realdokumenten (V7/V8/V10/
# V11, verschiedene Autoren) zurechtkommen und sucht Felder daher per
# Text-Heuristik. Der Filler bearbeitet dagegen IMMER dieselbe, fest
# bekannte Leer-Vorlage - deshalb sind hier feste Tabellen-/Zeilen-/
# Zellen-Indizes verwendet (robuster und einfacher als Text-Suche),
# ermittelt durch direkte Inspektion von
# assets/templates_docx/Systembewertung_V11_leer.docx.
#
# TEMPLATE-VERSION FEST VERANKERT: assets/templates_docx/
# Systembewertung_V11_leer.docx ist DIE im System hinterlegte,
# maßgebliche Vorlage fuer alle neu erzeugten Systembewertungen -
# bewusst als normale Datei im Repo abgelegt (sichtbar/ersetzbar bei
# Bedarf, kein verstecktes Binaerformat), aber NICHT einfach durch
# eine andere Datei ersetzbar: fill_template() prueft beim Laden per
# common.detect_template_version(), dass die Datei tatsaechlich V11
# ist, und bricht mit einer klaren Fehlermeldung ab, falls nicht.
#
# Ein Wechsel auf eine neue Template-Version (z.B. V12) darf NICHT
# durch einfaches Austauschen dieser Datei erfolgen - die Tabellen-/
# Zeilen-/Zellen-Indizes in diesem Modul (fill_deckblatt_rollen,
# fill_klassifizierung, CHECKBOX_MAPPING_V11 usw.) sind exakt auf die
# Struktur DIESER Datei abgestimmt und wuerden bei einer strukturell
# abweichenden Vorlage falsche Zellen befuellen, ohne dass das
# auffaellt. Eine neue Version muss daher ueber eine echte
# Code-Anpassung (Struktur-Analyse + Anpassung der Fill-Funktionen +
# Round-Trip-Test gegen die Extraktion, wie bei V11 geschehen) erfolgen
# - dafuer wieder Claude Code hinzuziehen, nicht die Datei manuell
# ersetzen.
#
# BEKANNTE EINSCHRAENKUNG (siehe auch README.md): Kapitel 5 (der
# Entscheidungsbaum Geraetekategorie/CS-Typ selbst, inkl. der
# zugehoerigen Ja/Nein-Antworten, Tabellen 9-10) wird NICHT automatisch
# befuellt, weil sich aus dem in der Zusammenfassungstabelle
# gespeicherten Endergebnis nicht eindeutig der zugrunde liegende
# Entscheidungsweg rekonstruieren laesst (mehrere Antwortpfade koennen
# zum selben Endergebnis fuehren - ein Raten waere in einem GxP-
# Dokument nicht vertretbar). Kapitel 5 muss nach der automatischen
# Erstellung manuell in Word ergaenzt werden.
#
# Kapitel 3 (Systemeinstufung Globales CS), 6 (ERES-Typ) und 7 (GAMP5-
# Kategorie) fragen NUR dieselbe Information nochmal ab, die auch in
# der Zusammenfassungstabelle (Kapitel 2) steht - dort gibt es klare
# 1:1-Checkboxen ohne Entscheidungsbaum, deshalb werden sie mit
# denselben Werten automatisch mitbefuellt (fill_kapitel3/6/7). Die
# Testtiefe (Kapitel 2 direkt + die Z-Felder-Matrix in Kapitel 8) wird
# aus GxP-Kritikalitaet + GAMP5 Software-Kategorie automatisch berechnet
# (fill_testtiefe) - ebenfalls kein Entscheidungsbaum, sondern eine
# feste Matrix. Alles, was in der Zusammenfassungstabelle (Kapitel 2)
# und auf dem Deckblatt steht, wird vollstaendig befuellt.
# ============================================================

import os
import re
from datetime import date

from docx import Document

from sysbew_common import get_cell_text

_NS = {
    "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
}

# Im System fest verankerte Vorlage - siehe Modul-Kommentar oben.
# TEMPLATE_VERSION ist die einzige Stelle, die bei einer neuen
# Template-Generation angepasst werden muss (nach Umsetzung der
# zugehoerigen Fill-Funktionen fuer die neue Struktur).
#
# Es gibt zu jedem Zeitpunkt IMMER NUR GENAU EINE aktive Version: kein
# Auswahl-Dropdown, kein Fallback auf eine aeltere Version. Steht z.B.
# V12 zur Verfuegung, wird TEMPLATE_VERSION auf 12 und
# DEFAULT_TEMPLATE_PATH auf die neue Datei umgestellt, die alte
# V11-Datei wird entfernt (nicht parallel weiter angeboten) - die
# Umstellung erfolgt ausschliesslich ueber eine Code-Aenderung
# (Claude Code hinzuziehen), niemals durch Ersetzen der Datei allein.
TEMPLATE_VERSION = 11
DEFAULT_TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "assets", "templates_docx", "Systembewertung_V11_leer.docx",
)

class TemplateVersionError(RuntimeError):
    """Die unter DEFAULT_TEMPLATE_PATH/template_path hinterlegte Datei
    ist keine gueltige V11-Vorlage (falsche Version, versehentlich
    ersetzt, oder beschaedigt)."""

# ============================================================
# Low-Level: Checkbox-Zustand schreiben (inkl. sichtbarem Glyph)
# ============================================================
def _w14(tag):
    return f"{{{_NS['w14']}}}{tag}"

def _w(tag):
    return f"{{{_NS['w']}}}{tag}"

def get_checkbox_sdts_from_cell(cell):
    """Alle echten Checkbox-SDTs einer Zelle, in Dokumentreihenfolge
    (dieselbe Reihenfolge, die sysbew_common.get_checkboxes_from_cell
    beim Lesen verwendet - wichtig, damit cb_idx aus CHECKBOX_MAPPING_V11
    zur richtigen Checkbox passt)."""
    result = []
    for sdt in cell._element.findall(".//w:sdt", _NS):
        checkbox = sdt.find(".//w14:checkbox", _NS)
        if checkbox is not None:
            result.append(sdt)
    return result

def set_checkbox_state_in_sdt(sdt, checked):
    """Setzt sowohl den logischen Zustand (w14:checked/@w14:val) als
    auch das sichtbare Glyphen-Zeichen (☐/☒) im sdtContent - beides
    muss uebereinstimmen, sonst zeigt Word beim Oeffnen den falschen
    Haken an, bis man die Checkbox einmal anklickt."""
    checkbox = sdt.find(".//w14:checkbox", _NS)
    if checkbox is None:
        return False
    checked_el = checkbox.find("w14:checked", _NS)
    if checked_el is None:
        checked_el = checkbox.makeelement(_w14("checked"), {})
        checkbox.insert(0, checked_el)
    checked_el.set(_w14("val"), "1" if checked else "0")

    state_tag = "checkedState" if checked else "uncheckedState"
    state_el = checkbox.find(f"w14:{state_tag}", _NS)
    glyph = None
    if state_el is not None:
        code_hex = state_el.get(_w14("val"))
        if code_hex:
            glyph = chr(int(code_hex, 16))

    if glyph:
        sdt_content = sdt.find("w:sdtContent", _NS)
        if sdt_content is not None:
            t = sdt_content.find(".//w:t", _NS)
            if t is not None:
                t.text = glyph
    return True

def set_checkbox_in_paragraph(paragraph, checked):
    """Wie set_checkboxes_in_cell(), aber fuer eine Checkbox, die in
    einem eigenstaendigen Absatz steht statt in einer Tabellenzelle
    (z.B. das "N/A - Weiter mit Kap. X"-Muster, das mehrfach im
    Template vorkommt, ausserhalb jeder Tabelle). Setzt nur die ERSTE
    Checkbox im Absatz - fuer die bekannten Faelle reicht das, da dort
    jeweils genau eine steht."""
    for sdt in paragraph._p.findall(".//w:sdt", _NS):
        if sdt.find(".//w14:checkbox", _NS) is not None:
            set_checkbox_state_in_sdt(sdt, checked)
            return True
    return False

def set_checkboxes_in_cell(cell, states):
    """`states`: dict {checkbox_index: bool}. Checkbox-Indizes, die
    nicht im dict stehen, bleiben unveraendert (Default des leeren
    Templates = nicht angekreuzt)."""
    sdts = get_checkbox_sdts_from_cell(cell)
    for idx, checked in states.items():
        if idx < len(sdts):
            set_checkbox_state_in_sdt(sdts[idx], bool(checked))

def set_cell_text(cell, text):
    """Ersetzt den gesamten Zellinhalt durch `text` (mehrzeilig via
    "\\n" wird weiterhin als Zeilenumbruch dargestellt). Leerer/None-
    Text wird als leere Zelle geschrieben (Platzhalter wird also auch
    dann entfernt, wenn keine Daten da sind).

    Schreibt bewusst NICHT ueber `cell.text = text` (python-docx loescht
    dabei die komplette Zelle und legt einen NEUEN Absatz/Run mit
    Standardformatierung an) - das hat im Template-Platzhaltertext
    (z.B. "<<Vorname Nachname>>") eine andere Schriftart/-groesse als
    der Rest des Dokuments zur Folge. Stattdessen wird der Text in den
    ERSTEN bestehenden Run des ersten Absatzes geschrieben (dessen
    Formatierung bleibt erhalten); alle weiteren Runs der Zelle werden
    geleert und zusaetzliche Absaetze anschliessend entfernt (siehe
    unten).

    Leert dabei bewusst ALLE <w:t>-Textknoten der Zelle direkt per XML
    (nicht nur ueber Paragraph.runs) - Bugfix: manche Platzhaltertexte
    enthalten einen <w:hyperlink> (z.B. ein Verweis auf "QU-MT-0001344"
    in der Dokumentenhistorie-Zeile), dessen Run python-docx's
    Paragraph.runs NICHT erfasst (nur direkte <w:r>-Kindelemente eines
    Absatzes, keine ueber <w:hyperlink> verschachtelten). Ohne diesen
    Direktzugriff bliebe so ein Hyperlink-Textfragment auch nach dem
    "Leeren" der Zelle stehen.

    Entfernt zusaetzliche Absaetze (2. Zeile und weiter) ANSCHLIESSEND
    komplett aus der Zelle, statt sie nur zu leeren - Bugfix: mehrere
    Platzhaltertexte im Leer-Template sind selbst zweizeilig (z.B.
    "Bezeichnung des Equipments/Systemname:" oder "Hersteller/
    SW-Ersteller/Lieferant:"), was sonst nach dem Befuellen eine leere
    Zeile im fertigen Dokument haette stehen lassen. Mehrzeilige WERTE
    ("\\n" im uebergebenen `text`) sind davon nicht betroffen - die
    werden ueber Run.text als Zeilenumbruch INNERHALB des ersten Laufs
    dargestellt, nicht als zusaetzlicher Absatz.

    Entfernt ausserdem <w:br/>/<w:tab/>-Elemente in ANDEREN Laeufen
    desselben Absatzes komplett - Bugfix: mehrzeilige Platzhalter-
    Anleitungstexte (z.B. bei "Parameter:") haben ihren Zeilenumbruch
    oft MITTEN in einem spaeteren Lauf des ERSTEN Absatzes stehen
    (<w:br/> + <w:t> im selben <w:r>) statt in einem eigenen Absatz -
    das blosse Leeren des zugehoerigen <w:t> liess den <w:br/> selbst
    stehen und erzeugte dadurch eine leere Zeile bzw. einen
    ueberfluessigen Zeilenumbruch am Ende des neuen Werts."""
    text = text or ""
    absaetze = cell.paragraphs
    erster = absaetze[0] if absaetze else None
    if erster is None or not erster.runs:
        cell.text = text  # Fallback: keine Formatierung erkennbar/vorhanden
        return
    erster_lauf = erster.runs[0]
    erster_lauf.text = text
    # WICHTIG: `Run.text = "a\nb"` legt bei python-docx MEHRERE <w:t>-
    # Knoten (getrennt durch <w:br/>) INNERHALB desselben Laufs an -
    # deshalb hier den ganzen ersten Lauf schuetzen (nicht nur einen
    # einzelnen <w:t>-Knoten), sonst wuerde ein mehrzeiliger Text im
    # naechsten Schritt gleich wieder selbst geleert (Bugfix).
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for t in cell._element.findall(".//w:t", ns):
        if t.getparent() is not erster_lauf._r:
            t.text = ""
    for tag in ("w:br", "w:tab"):
        for el in cell._element.findall(f".//{tag}", ns):
            if el.getparent() is not erster_lauf._r:
                el.getparent().remove(el)
    for p in absaetze[1:]:
        p._p.getparent().remove(p._p)

def replace_in_cell_paragraphs(cell, platzhalter, ersatz):
    """Ersetzt `platzhalter` (Text-Teilstring, z.B. "(Site/Unit)") durch
    `ersatz` in jedem Absatz einer Zelle, der ihn enthaelt - OHNE den
    Rest des Absatzes anzutasten (anders als set_cell_text(), das die
    ganze Zelle ueberschreibt). Wird u.a. fuer die Abteilungs-Platzhalter
    in der Deckblatt-Rollentabelle gebraucht, wo der Platzhalter nur ein
    Teil eines laengeren Labeltexts ist (z.B. "TSO (Technical System
    Owner)  (Site/Unit)").

    Baut den gesamten (ggf. ueber mehrere Runs verteilten) Absatztext neu
    zusammen: der erste Run behaelt seine Formatierung und bekommt den
    ersetzten Text, alle weiteren Runs des Absatzes werden leer - das
    reicht fuer diese kurzen, einheitlich formatierten Label-Zeilen."""
    aendert = False
    for p in cell.paragraphs:
        voller_text = "".join(r.text for r in p.runs)
        if platzhalter not in voller_text:
            continue
        neuer_text = voller_text.replace(platzhalter, ersatz)
        if p.runs:
            p.runs[0].text = neuer_text
            for r in p.runs[1:]:
                r.text = ""
        else:
            p.text = neuer_text
        aendert = True
    return aendert

# ============================================================
# Checkbox-Gruppen: generisch ueber CHECKBOX_MAPPING_V11 (identische
# Struktur zu sysbew_common.extract_checkboxes_formularfelder, nur in
# Schreibrichtung)
# ============================================================
def fill_checkboxes_formularfelder(doc, checkbox_mapping, data):
    for table_idx, table_mapping in checkbox_mapping.items():
        table = doc.tables[table_idx]
        rows = list(table.rows)
        for row_idx, row_mapping in table_mapping.items():
            cells = rows[row_idx].cells
            for cell_idx, cell_mapping in row_mapping.items():
                cell = cells[cell_idx]
                states = {
                    cb_idx: data.get(col) == "r"
                    for cb_idx, col in cell_mapping.items()
                }
                set_checkboxes_in_cell(cell, states)

# ============================================================
# Deckblatt: Rollen/Namen
# ============================================================

# Abteilungs-Felder (webapp-only, nicht Teil von EXCEL_COLUMNS - siehe
# webapp/app.py, ABTEILUNG_FELDER): jede Rollenzeile im Template enthaelt
# im Label (Zelle 0) einen festen Abteilungs-Platzhalter, der bei
# Bedarf ersetzt wird - bei Ersteller/SI-PL/TSO/BSO/BQR "(Site/Unit)",
# bei CSQ das im Label fest eingetragene "(FBC Quality Q&V CSV)" (der
# Wert wird nur ersetzt, wenn tatsaechlich ein CSQ_Abteilung-Wert
# angegeben wird - Standardfall bleibt unangetastet). SME hat keine
# eigene Zeile (siehe unten).
_ABTEILUNG_FELD_JE_ROLLE = {
    "Ersteller": "Ersteller_Abteilung",
    "SI/PL":     "SI_PL_Abteilung",
    "TSO":       "TSO_Abteilung",
    "BSO":       "BSO_Abteilung",
    "BQR":       "BQR_Abteilung",
    # CSQ hat KEINEN "(Site/Unit)"-Platzhalter wie die anderen Rollen,
    # sondern ein im Label fest eingetragenes "(FBC Quality Q&V CSV)" -
    # dieser Text wird bei Bedarf ersetzt (siehe
    # _ABTEILUNG_PLATZHALTER_JE_ROLLE unten), bleibt aber unangetastet,
    # wenn kein CSQ_Abteilung-Wert angegeben ist (Standardfall).
    "CSQ":       "CSQ_Abteilung",
}
_ABTEILUNG_PLATZHALTER_JE_ROLLE = {
    "Ersteller": "(Site/Unit)",
    "SI/PL":     "(Site/Unit)",
    "TSO":       "(Site/Unit)",
    "BSO":       "(Site/Unit)",
    "BQR":       "(Site/Unit)",
    "CSQ":       "(FBC Quality Q&V CSV)",
}

def fill_deckblatt_rollen(doc, data):
    """Tabelle 0 (Unterschriftentabelle): Zeile 0/2/4/6/8/10, Spalte 1
    (siehe ROLLEN_SPALTEN-Reihenfolge in sysbew_common.py)."""
    table = doc.tables[0]
    zeilen_je_rolle = {
        "Ersteller": 0, "SI/PL": 2, "TSO": 4,
        "BSO": 6, "BQR": 8, "CSQ": 10,
    }
    for spalte, row_idx in zeilen_je_rolle.items():
        wert = data.get(spalte, "")
        # set_cell_text() bewusst IMMER aufrufen (siehe Bugfix-
        # Kommentar in fill_kapitel1) - das Leer-Template hat hier den
        # Platzhalter "<<Vorname Nachname>>" statt einer leeren Zelle.
        set_cell_text(table.rows[row_idx].cells[1], wert)
        abteilung_feld = _ABTEILUNG_FELD_JE_ROLLE.get(spalte)
        if abteilung_feld:
            abteilung = data.get(abteilung_feld, "")
            if abteilung:
                platzhalter = _ABTEILUNG_PLATZHALTER_JE_ROLLE[spalte]
                replace_in_cell_paragraphs(
                    table.rows[row_idx].cells[0],
                    platzhalter, f"({abteilung})",
                )
    # SME: nur befuellen, wenn nicht ohnehin schon mit SI/PL identisch
    # (kombinierte Rolle "Projektleiter/SME" - Name steht dann bereits
    # in Zeile 2 bei SI/PL). Ansonsten SME wie eigene Rolle behandeln;
    # das Template hat dafuer keine eigene Zeile, SME wird deshalb an
    # die SI/PL-Zeile angehaengt.
    #
    # WICHTIG: der bestehende Zellentext wird hier NICHT zurueckgelesen
    # (Bugfix) - ist SI/PL leer, wurde die Zeile oben NICHT beschrieben
    # (der "if wert:"-Schutz greift), der Platzhalter "<<Vorname
    # Nachname>>" stuende dann noch drin und wuerde faelschlich vor den
    # SME-Namen gehaengt. Stattdessen wird der bekannte Datenwert
    # `si_pl` verwendet - leer, wenn keine SI/PL-Person angegeben ist.
    sme = data.get("SME", "")
    si_pl = data.get("SI/PL", "")
    if sme and sme != si_pl:
        neu = f"{si_pl}\n{sme}" if si_pl else sme
        set_cell_text(table.rows[2].cells[1], neu)

# ============================================================
# Klassifizierung (Tabelle 3, Zeile 1)
# ============================================================
def fill_klassifizierung(doc, data):
    row = doc.tables[3].rows[1]
    set_checkboxes_in_cell(row.cells[0], {0: data.get("KLASS_Lokal") == "r"})
    set_checkboxes_in_cell(row.cells[1], {
        0: data.get("KLASS_Multisite") == "r",
        1: data.get("KLASS_Multisite_NurLokal") == "r",
        2: data.get("KLASS_Multisite_LokalGlobal") == "r",
    })
    set_checkboxes_in_cell(row.cells[2], {
        0: data.get("KLASS_Global") == "r",
        1: data.get("KLASS_Global_1a") == "r",
        2: data.get("KLASS_Global_1b") == "r",
        3: data.get("KLASS_Global_2") == "r",
        4: data.get("KLASS_Global_3") == "r",
    })
    set_checkboxes_in_cell(row.cells[3], {0: data.get("KLASS_OhneCS") == "r"})

# ============================================================
# Kapitel 1 (Tabelle 2): Neuerstellung, Equipment-Name,
# Kurzbeschreibung, Systemtyp-Zugang, Einsatzbereich, Hersteller
# ============================================================
def fill_kapitel1(doc, data):
    table = doc.tables[2]

    set_checkboxes_in_cell(table.rows[1].cells[0], {
        0: data.get("Neuerstellung") == "r",
        1: data.get("Revisioniert") == "r",
    })
    # Grund der Systembewertung: ersetzt den Hinweistext "(Grund: Gemäß
    # CC-Nummer und ggf. Text/PR) (CC-Nr. und ggf. Text/PR/Re-
    # Qualifizierung/Re-Validierung)" durch den bereits an anderer
    # Stelle erfassten Grund (Historie/"Grund der Erstellung", siehe
    # auch fill_historie() fuer die Dokumentenhistorie-Tabelle - hier
    # dieselbe Angabe direkt neben Neuerstellung/Änderung).
    # set_cell_text() bewusst IMMER aufrufen, auch bei leerem Feld -
    # das Leer-Template hat in diesen Zellen keinen Blanko-Platz,
    # sondern ausgeschriebene Anleitungs-/Beispieltexte ("Bei
    # Equipment z.B: Laborwaage...", "(Grund: Gemäß CC-Nummer...)"
    # usw.). Ein "nur schreiben, wenn befuellt" wuerde diese Anleitung
    # bei fehlender Angabe unbemerkt im fertigen Dokument stehen
    # lassen, statt die Zelle wie vorgesehen leer zu lassen (Bugfix -
    # siehe README.md, Versionshistorie).
    set_cell_text(table.rows[1].cells[1], data.get("Historie") or "")
    set_cell_text(table.rows[3].cells[1], data.get("AS/BDIS-Name") or "")
    set_cell_text(table.rows[4].cells[1], data.get("Kurzbeschreibung") or "")

    set_checkboxes_in_cell(table.rows[5].cells[1], {
        0: data.get("Offen") == "r",
        1: data.get("Geschlossen") == "r",
        2: data.get("NA") == "r",
    })
    # Begruendung zum Systemtyp (webapp-only, nicht Teil von
    # EXCEL_COLUMNS - siehe webapp/app.py): das Template hat dafuer
    # keinen eigenen Platzhalter (nur bei "N/A" steht der feste
    # Beispieltext "Begründung: mechanische Ausrüstung" dabei). Wird
    # als NEUER Absatz angehaengt statt per set_cell_text() die Zelle
    # zu ueberschreiben - die Zelle enthaelt 3 Absaetze mit je einer
    # eigenen Checkbox (SDT) + Label, set_cell_text() wuerde deren
    # Labeltexte zerstoeren (nur die Checkbox-SDTs selbst blieben
    # erhalten, siehe set_cell_text()-Dokumentation).
    if data.get("SystemtypZugang_Begruendung"):
        table.rows[5].cells[1].add_paragraph(
            f"Begründung: {data['SystemtypZugang_Begruendung']}"
        )

    einsatzbereich = data.get("Betrieb", "")
    if data.get("Gebaeude"):
        einsatzbereich = f"{einsatzbereich}/{data['Gebaeude']}" if einsatzbereich else data["Gebaeude"]
    set_cell_text(table.rows[6].cells[1], einsatzbereich)

    hersteller = data.get("Hersteller", "")
    if data.get("SW-Hersteller") and data["SW-Hersteller"] != hersteller:
        hersteller = f"{hersteller} / {data['SW-Hersteller']}" if hersteller else data["SW-Hersteller"]
    # Lieferantennummer (QualiPSO-/QTP-Customer-ID) hat keine eigene
    # Zelle im Template - gehoert inhaltlich zum Hersteller-Feld und
    # wird deshalb dort angehaengt (wie beim Lesen echter Dokumente,
    # wo dieselbe Angabe umgekehrt aus genau diesem Text herausgelesen
    # wird, siehe parse_hersteller() in word_parser_v8/10/11.py).
    if data.get("Lieferantennummer"):
        zusatz = f"QualiPSO-ID: {data['Lieferantennummer']}"
        hersteller = f"{hersteller} / {zusatz}" if hersteller else zusatz
    set_cell_text(table.rows[7].cells[1], hersteller)

# ============================================================
# Business Kritikalitaet (Tabelle 4) + GxP-Relevanz (Tabelle 5)
# - Detailfragen, deren Ergebnis in der Zusammenfassungstabelle
# (Kapitel 2) noch einmal auftaucht (dort ueber
# fill_checkboxes_formularfelder/CHECKBOX_MAPPING_V11 abgedeckt).
# ============================================================
def fill_business_kritikalitaet(doc, data):
    table = doc.tables[4]
    set_checkboxes_in_cell(table.rows[1].cells[0], {0: data.get("BCkritisch") == "r"})
    set_checkboxes_in_cell(table.rows[2].cells[0], {0: data.get("BCunkritisch") == "r"})

def fill_gxp_relevanz(doc, data):
    table = doc.tables[5]
    set_checkboxes_in_cell(table.rows[1].cells[0], {0: data.get("GxP_Relevan_JA") == "r"})
    set_checkboxes_in_cell(table.rows[2].cells[0], {0: data.get("GxP_Relevan_NEIN") == "r"})

# ============================================================
# Zusammenfassungstabelle (Tabelle 6): MLCS-ID, Anlage,
# Schnittstelle/UeberlagerteMLCS, Besonderheiten
# (die Checkbox-Spalten r6/r8 laufen ueber CHECKBOX_MAPPING_V11)
# ============================================================
def fill_zusammenfassung_text(doc, data):
    table = doc.tables[6]
    # set_cell_text() bewusst IMMER aufrufen (siehe Bugfix-Kommentar in
    # fill_kapitel1) - auch diese Zellen haben im Leer-Template
    # ausgeschriebene Anleitungstexte statt einer leeren Zelle.
    set_cell_text(table.rows[1].cells[3], data.get("MLCSID") or "")
    set_cell_text(table.rows[2].cells[3], data.get("Anlage") or "")

    schnittstelle = data.get("Schnittstelle") or data.get("UeberlagerteMLCS") or ""
    if data.get("Schnittstelle") and data.get("UeberlagerteMLCS") and \
            data["Schnittstelle"] != data["UeberlagerteMLCS"]:
        schnittstelle = f"{data['Schnittstelle']} / {data['UeberlagerteMLCS']}"
    set_cell_text(table.rows[3].cells[3], schnittstelle)

    if data.get("Besonderheiten"):
        # Nicht ersetzen, sondern ergaenzen: die Zelle enthaelt im
        # Leer-Template einen Hinweistext ("Bei Gerätekategorien A, B
        # und C bitte die Subkategorisierung ... nach QU-SOP-0021736
        # angeben") - der bleibt als Anleitung stehen, die eigentliche
        # Besonderheiten-Angabe wird dahinter angehaengt.
        zelle = table.rows[9].cells[1]
        bestehend = zelle.text.strip()
        neu = f"{bestehend}\n{data['Besonderheiten']}" if bestehend else data["Besonderheiten"]
        set_cell_text(zelle, neu)

    # Periodic Review "andere/freie Angabe": die Checkbox selbst laeuft
    # ueber CHECKBOX_MAPPING_V11 (PR_Andere), hier wird zusaetzlich der
    # dahinterstehende Blanko-Platzhalter ("_______________") durch den
    # tatsaechlichen Freitext ersetzt, falls angegeben (webapp-only
    # Zusatzfeld "PR_Andere_Text", nicht Teil von EXCEL_COLUMNS).
    if data.get("PR_Andere") == "r" and data.get("PR_Andere_Text"):
        replace_in_cell_paragraphs(
            table.rows[4].cells[3], "_______________", data["PR_Andere_Text"],
        )

# ============================================================
# Kapitel 3 (Tabelle 7): Systemeinstufung Globales CS - Detail-
# Entscheidung, ob Klasse 1a oder 1b vorliegt (Nein/Ja-Frage), plus
# Klasse 2/Klasse 3. Fachlich dieselbe Information wie die Klasse-1a/
# 1b/2/3-Checkboxen der Klassifizierung in Kapitel 1 (Tabelle 3) - es
# gibt dafuer bewusst KEINE eigenen Excel-Spalten, sondern es werden
# dieselben KLASS_Global_1a/1b/2/3-Werte hier zusaetzlich eingetragen,
# damit Kapitel 1 und Kapitel 3 im erzeugten Dokument konsistent sind
# (vorher wurde Kapitel 3 gar nicht befuellt).
# ============================================================
def fill_kapitel3(doc, data):
    table = doc.tables[7]
    set_checkboxes_in_cell(table.rows[0].cells[0], {
        0: data.get("KLASS_Global_1a") == "r",  # "Nein" -> Klasse 1a
        1: data.get("KLASS_Global_1b") == "r",  # "Ja"   -> Klasse 1b
    })
    set_checkboxes_in_cell(table.rows[1].cells[2], {0: data.get("KLASS_Global_2") == "r"})
    set_checkboxes_in_cell(table.rows[1].cells[3], {0: data.get("KLASS_Global_3") == "r"})

    # "N/A - Weiter mit Kap. 4": eigener Absatz VOR der obigen Tabelle
    # (kein Tabellen-Feld, siehe set_checkbox_in_paragraph) - Kapitel 3
    # fragt nur die Detailfrage INNERHALB von "Globales CS" ab (welche
    # Klasse 1a/1b/2/3), ist also nicht relevant, wenn das System gar
    # nicht als "Globales CS" klassifiziert ist (Kapitel 1 -
    # Klassifizierung: Lokales CS/Multi-Site-CS/Equipment ohne CS statt
    # Globales CS). Absatz-Index 19 (doc.paragraphs) per direkter
    # Inspektion des Templates ermittelt - der Absatz-Text lautet
    # ausschliesslich "N/A", direkt gefolgt vom Absatz "Weiter mit
    # Kap. 4".
    if data.get("KLASS_Global") != "r":
        set_checkbox_in_paragraph(doc.paragraphs[19], True)

# ============================================================
# GxP-Risikoklassifizierung im Detail (Tabelle 8) - Duplikat der
# GxP-Kritikalitaet aus der Zusammenfassungstabelle plus Begruendung
# ============================================================
def fill_gxp_risikoklassifizierung(doc, data):
    table = doc.tables[8]
    set_checkboxes_in_cell(table.rows[1].cells[0], {
        0: data.get("GxP-C") == "r",
        1: data.get("GxP-M") == "r",
        2: data.get("GxP-m2") == "r",
        3: data.get("GxP-NA") == "r",
    })

    teile = []
    for label, feld in [
        ("Produktqualität", "GxP_Produktqualitaet"),
        ("Patientensicherheit", "GxP_Patientensicherheit"),
        ("Datenintegrität", "GxP_Datenintegritaet"),
    ]:
        wert = data.get(feld, "")
        if wert:
            teile.append(f"{label}: {wert}" if not wert.lower().startswith(label.lower()) else wert)
    begruendung = " ".join(teile).strip()
    # set_cell_text() bewusst IMMER aufrufen (siehe Bugfix-Kommentar in
    # fill_kapitel1) - das Leer-Template hat hier eine ausfuehrliche
    # Ausfuell-Anleitung statt einer leeren Zelle.
    set_cell_text(
        table.rows[2].cells[0],
        f"Begründung der GxP Risikoklassifizierung: {begruendung}" if begruendung else "",
    )

# ============================================================
# Kapitel 6 (Tabelle 11): ERES-Typ im Detail - dieselbe Auswahl wie
# ERESTYP1-4/NA in der Zusammenfassungstabelle (Kapitel 2), hier
# zusaetzlich an der Kapitel-6-Frage selbst eingetragen. Bei ERES Typ 4
# zusaetzlich "Art der Signatur" (3 webapp-only Zusatz-Checkboxen, nicht
# Teil von EXCEL_COLUMNS - im Template als eigene Unterfrage nur bei
# Typ 4 vorhanden).
# ============================================================
def fill_kapitel6(doc, data):
    table = doc.tables[11]
    zeilen = {
        "ERESTYP1": 1, "ERESTYP2": 2, "ERESTYP3": 3,
        "ERESTYP4": 4, "ERESTYPNA": 5,
    }
    for feld, row_idx in zeilen.items():
        if data.get(feld) == "r":
            set_checkboxes_in_cell(table.rows[row_idx].cells[1], {0: True})

    if data.get("ERESTYP4") == "r":
        set_checkboxes_in_cell(table.rows[4].cells[2], {
            0: data.get("ERES4_SIG_ID_PW") == "r",
            1: data.get("ERES4_SIG_BIOMETRISCH") == "r",
            2: data.get("ERES4_SIG_TOKEN_PW") == "r",
        })

# ============================================================
# Kapitel 7 (Tabelle 12): GAMP5-Software-Kategorie im Detail -
# dieselbe Auswahl wie GAMP5 Software-Kategorie in der
# Zusammenfassungstabelle (Kapitel 2), hier zusaetzlich an der
# Kapitel-7-Frage selbst eingetragen (kein N/A-Fall vorgesehen).
# ============================================================
def fill_kapitel7(doc, data):
    table = doc.tables[12]
    zeilen = {"KAT1": 1, "KAT3": 2, "KAT4": 3, "KAT5": 4}
    for feld, row_idx in zeilen.items():
        if data.get(feld) == "r":
            set_checkboxes_in_cell(table.rows[row_idx].cells[1], {0: True})

# ============================================================
# Kapitel 9 (Tabelle 14): 9.1 "Kommt KI zum Einsatz?" - ergibt sich
# direkt aus der KI-Reifegrad-Auswahl der Zusammenfassungstabelle
# (Kapitel 2): KINA (N/A) bedeutet kein KI-Einsatz -> "Nein", jede
# andere Reifegrad-Auswahl (I-VI) bedeutet KI-Einsatz -> "Ja". Seit dem
# Wegfall von "KI-Reifegrad" im Web-Editor (siehe SKIP_FELDER in
# webapp/app.py - die genaue Stufe laesst sich ohne die Detailfragen
# 9.2-9.5 nicht zuverlaessig abfragen) kommt die Antwort dort direkt
# ueber die einfachere Ja/Nein-Frage "Kommt KI zum Einsatz?"
# (webapp-only KI_Einsatz_Ja/KI_Einsatz_Nein) - wird hier zusaetzlich
# zu KINA/KI1-6 ausgewertet, damit 9.1 auch bei "Ja" ohne konkrete
# Stufe korrekt "Ja" zeigt. Die weiteren Detailfragen 9.2-9.5
# (verbotene Praktiken, Autonomie-Stufe, Steuerungsdesign-Stufe)
# haengen von zusaetzlichen Antworten ab, die daraus NICHT
# rekonstruierbar sind - werden daher bewusst nicht befuellt (wie
# Kapitel 5, siehe Modul-Kommentar oben).
# ============================================================
def fill_kapitel9(doc, data):
    table = doc.tables[14]
    if data.get("KINA") == "r" or data.get("KI_Einsatz_Nein") == "r":
        set_checkboxes_in_cell(table.rows[2].cells[1], {0: True})
    elif data.get("KI_Einsatz_Ja") == "r" or any(
        data.get(f) == "r" for f in ("KI1", "KI2", "KI3", "KI4", "KI5", "KI6")
    ):
        set_checkboxes_in_cell(table.rows[1].cells[1], {0: True})

# ============================================================
# Testtiefe (Zusammenfassungstabelle Kapitel 2, Tabelle 6 Zeile 6
# Zelle 4 UND Z-Felder-Matrix in Kapitel 8): wird automatisch aus GxP-
# Kritikalitaet (Kapitel 2) und GAMP5 Software-Kategorie (Kapitel 2)
# berechnet - dieselbe Matrix, die word_parser_v11.berechne_testtiefe()
# beim LESEN eines echten Dokuments aus den Z-Feldern zurueckrechnet,
# hier in Schreibrichtung. Keine manuelle Nacharbeit mehr noetig, sofern
# GxP-Kritikalitaet und Software-Kategorie ausgefuellt sind (bei N/A
# bei einem der beiden bleibt die Testtiefe bewusst leer).
# ============================================================
_TESTTIEFE_ZEILE_JE_KRITIKALITAET = {"GxP-C": "critical", "GxP-M": "major", "GxP-m2": "minor"}
_TESTTIEFE_SPALTE_JE_KATEGORIE = {"KAT1": 0, "KAT3": 0, "KAT4": 1, "KAT5": 2}
_TESTTIEFE_Z_FELDER = {
    "critical": ("Z1S1", "Z2S1", "Z3S1"),
    "major":    ("Z1S2", "Z2S2", "Z3S2"),
    "minor":    ("Z1S3", "Z2S3", "Z3S3"),
}

def fill_testtiefe(doc, data):
    kritikalitaet_feld = next(
        (f for f in _TESTTIEFE_ZEILE_JE_KRITIKALITAET if data.get(f) == "r"), None,
    )
    kategorie_feld = next(
        (f for f in _TESTTIEFE_SPALTE_JE_KATEGORIE if data.get(f) == "r"), None,
    )
    if not kritikalitaet_feld or not kategorie_feld:
        return  # GxP-NA oder SW-Kat-NA -> keine Testtiefe ableitbar

    zeile = _TESTTIEFE_ZEILE_JE_KRITIKALITAET[kritikalitaet_feld]
    spalte = _TESTTIEFE_SPALTE_JE_KATEGORIE[kategorie_feld]
    z_feld = _TESTTIEFE_Z_FELDER[zeile][spalte]

    if z_feld in ("Z2S1", "Z3S1", "Z3S2"):
        ttiefe = "TTIEFEHOCH"
    elif z_feld in ("Z1S1", "Z2S2", "Z3S3"):
        ttiefe = "TTIEFEMITTEL"
    else:
        ttiefe = "TTIEFENIEDRIG"

    # 1) Zusammenfassungstabelle Kapitel 2 - dieselbe Zelle, die V8/V10
    #    direkt per Checkbox abfragen (bei V11 im Template zwar
    #    vorhanden, wird laut Praxis aber meist nur ueber Kapitel 8
    #    bestimmt - wir befuellen sie trotzdem konsistent mit).
    set_checkboxes_in_cell(doc.tables[6].rows[6].cells[4], {
        0: ttiefe == "TTIEFENIEDRIG",
        1: ttiefe == "TTIEFEMITTEL",
        2: ttiefe == "TTIEFEHOCH",
    })

    # 2) Z-Felder-Matrix in Kapitel 8 (nested Tabelle in Tabelle 13)
    for table in doc.tables:
        table_text = " ".join(
            get_cell_text(c) for row in table.rows for c in row.cells
        ).lower()
        if "8.0" not in table_text and "festlegung der testtiefe" not in table_text:
            continue
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    nested_text = " ".join(
                        get_cell_text(c) for r in nested.rows for c in r.cells
                    ).lower()
                    if not all(k in nested_text for k in ("critical", "major", "minor")):
                        continue
                    for nrow in nested.rows:
                        label = get_cell_text(nrow.cells[0]).strip().lower()
                        if label != zeile:
                            continue
                        for i in range(3):
                            set_checkboxes_in_cell(nrow.cells[1 + i], {0: i == spalte})
                        return

# ============================================================
# Historie (Tabelle 1, Zeile 1) - fuer eine NEUE Systembewertung
# immer Version 1.0 mit heutigem Datum
# ============================================================
def fill_historie(doc, data):
    table = doc.tables[1]
    version = data.get("Version_Historie") or "1.0"
    datum = data.get("Datum") or date.today().strftime("%d.%m.%Y")
    set_cell_text(table.rows[1].cells[0], f"{version} / {datum}")

    grund = data.get("Historie") or (
        "Neuerstellung" if data.get("Neuerstellung") == "r" else ""
    )
    # set_cell_text() bewusst IMMER aufrufen (siehe Bugfix-Kommentar in
    # fill_kapitel1) - das Leer-Template hat hier den Beispieltext
    # "Neuerstellung auf Basis der QU-MT-0001344 Version 11.0, ...",
    # der ohne Angabe faelschlich stehen bliebe (und bei einer
    # tatsaechlichen Aenderung/Revision sogar irrefuehrend waere).
    set_cell_text(table.rows[1].cells[1], grund)

# ============================================================
# Beschreibungstabelle (Tabelle 16)
# ============================================================
_BESCHREIBUNG_ZEILEN = {
    0: "Prozessbeschreibung",
    1: "Daten",
    2: "Parameter",
    3: "Alarme (GxP-relevant)",
    4: "Chargenprotokoll",
    5: "Audit Trail (AT)",
    6: "Benutzer-verwaltung?",
    7: "Schnittstellen mit PLS",
    # Zeile 8 ("Datenfluss / Abbildung:") - webapp-only Feld
    # "DatenflussAbbildung" (nicht Teil von EXCEL_COLUMNS, siehe
    # webapp/app.py), da diese Zeile im Template primaer eine Grafik
    # erwartet und in der Master-Excel nie eine eigene Spalte hatte.
    8: "DatenflussAbbildung",
    9: "Angeschlossenes Equipment",
    10: "Sonstiges",
    11: "KI Bewertung",
}

# Die 4 generischen "BemerkungX"-Spalten der Master-Excel haben laut
# Fachbereich eine feste Bedeutung (siehe auch FELD_LABELS in
# webapp/app.py) und gehoeren inhaltlich zu den jeweiligen Zeilen
# dieser Tabelle ("Informationen und Bemerkungen") - werden also, falls
# befuellt, dort zusaetzlich eingetragen statt (wie bisher) gar nicht
# in ein erzeugtes Dokument uebernommen zu werden.
_BEMERKUNG_ZUORDNUNG = {
    "Bemerkung1": "Prozessbeschreibung",
    "Bemerkung2": "Daten",
    "Bemerkung3": "Audit Trail (AT)",
    "Bemerkung4": "Parameter",
}

def fill_beschreibungstabelle(doc, data):
    table = doc.tables[16]

    werte = {excel_col: data.get(excel_col) or "" for excel_col in _BESCHREIBUNG_ZEILEN.values()}

    # "Steuerung erfolgt über?" hat im Template KEINE eigene Zeile -
    # gehoert inhaltlich zur Prozessbeschreibung und wird deshalb dort
    # vorangestellt (vorher wurde dieser Wert erfasst, aber nie in ein
    # Dokument uebernommen).
    if data.get("Steuerung erfolgt über?"):
        vorspann = f"Steuerung erfolgt über: {data['Steuerung erfolgt über?']}"
        werte["Prozessbeschreibung"] = (
            f"{vorspann}\n{werte['Prozessbeschreibung']}" if werte["Prozessbeschreibung"] else vorspann
        )

    # BemerkungX anhaengen (nicht ersetzen, falls das dedizierte Feld
    # zusaetzlich befuellt wurde).
    for bemerkung_feld, excel_col in _BEMERKUNG_ZUORDNUNG.items():
        zusatz = data.get(bemerkung_feld)
        if zusatz:
            werte[excel_col] = f"{werte[excel_col]}\n{zusatz}" if werte[excel_col] else zusatz

    # set_cell_text() bewusst IMMER aufrufen, auch bei leerem Feld -
    # BUGFIX (siehe README.md, Versionshistorie): das Leer-Template hat
    # in JEDER dieser 12 Zeilen einen ausgeschriebenen Anleitungs-/
    # Fragetext in "<...>" (z.B. bei "Audit Trail / Audit Trail
    # Review:" die komplette, unbeantwortete Frage, bei "Angeschlossene
    # Equipments:" einen Beispielabsatz mit "<System>"/"<Firmware>").
    # Ein "nur schreiben, wenn befuellt" liess bislang genau diesen
    # Anleitungstext unbemerkt im fertigen, offiziellen Dokument stehen,
    # sobald ein Feld im Editor leer gelassen wurde - keine harmlose
    # Lücke, sondern falscher/unbeantworteter Inhalt in einer GxP-
    # Systembewertung.
    for row_idx, excel_col in _BESCHREIBUNG_ZEILEN.items():
        set_cell_text(table.rows[row_idx].cells[1], werte[excel_col])

# ============================================================
# HAUPTFUNKTION
# ============================================================
def fill_template(data, template_path=None, output_path=None):
    """Erzeugt aus `data` (dieselbe Feldstruktur wie EXCEL_COLUMNS/
    Master-Excel) eine neue Systembewertung auf Basis des V11-Leer-
    Templates. Gibt den docx.Document zurueck; speichert zusaetzlich
    unter `output_path`, falls angegeben.

    Importiert CHECKBOX_MAPPING_V11 hier (statt am Modulanfang), um
    einen Zirkelimport mit word_parser_v11.py (das umgekehrt aus
    sysbew_common importiert, nicht aus diesem Modul) zu vermeiden -
    ist unproblematisch, da fill_template() ohnehin erst zur Laufzeit
    aufgerufen wird.
    """
    from word_parser_v11 import CHECKBOX_MAPPING_V11
    import sysbew_common as common

    template_path = template_path or DEFAULT_TEMPLATE_PATH
    if not os.path.exists(template_path):
        raise TemplateVersionError(
            f"Template nicht gefunden: {template_path}\n"
            f"Die im System hinterlegte V{TEMPLATE_VERSION}-Vorlage fehlt oder "
            f"wurde verschoben - bitte assets/templates_docx/ pruefen."
        )
    doc = Document(template_path)

    erkannt = common.detect_template_version(doc)
    if erkannt != TEMPLATE_VERSION:
        raise TemplateVersionError(
            f"Das Template unter {template_path} ist keine gueltige "
            f"V{TEMPLATE_VERSION}-Vorlage (erkannt: V{erkannt}). Diese Datei "
            f"darf nicht einfach durch eine andere Version ersetzt werden - "
            f"die Tabellen-/Zellzuordnung in diesem Modul ist exakt auf die "
            f"V{TEMPLATE_VERSION}-Struktur abgestimmt. Eine neue Template-"
            f"Version muss ueber eine Code-Anpassung (Struktur-Analyse + "
            f"neue Fill-Funktionen + Round-Trip-Test) integriert werden."
        )

    fill_deckblatt_rollen(doc, data)
    fill_historie(doc, data)
    fill_kapitel1(doc, data)
    fill_klassifizierung(doc, data)
    fill_kapitel3(doc, data)
    fill_business_kritikalitaet(doc, data)
    fill_gxp_relevanz(doc, data)
    fill_zusammenfassung_text(doc, data)
    fill_checkboxes_formularfelder(doc, CHECKBOX_MAPPING_V11, data)
    fill_gxp_risikoklassifizierung(doc, data)
    fill_kapitel6(doc, data)
    fill_kapitel7(doc, data)
    fill_kapitel9(doc, data)
    fill_testtiefe(doc, data)
    fill_beschreibungstabelle(doc, data)

    if output_path:
        doc.save(output_path)
    return doc
