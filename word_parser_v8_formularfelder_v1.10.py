# ============================================================
# word_parser_v8_formularfelder.py
# V8 Systembewertung Extraktor - 1.10
# Drag & Drop: Word-Datei auf dieses Script ziehen!
# Output: Excel-Datei mit einer Zeile im gleichen Ordner
# ============================================================

from docx import Document
import re
import os
import sys
import time
from datetime import datetime

# ============================================================
# EXCEL-SPALTEN (exakt wie in der Vorlage, + 2 neue am Ende)
# ============================================================
EXCEL_COLUMNS = [
    "MLCSID", "API", "Betrieb", "Gebaeude", "Version", "Dok. -Nr.",
    "AS/BDIS-Name", "Anlage",
    "GxP_Relevan_JA", "GxP_Relevan_NEIN",
    "GxP-C", "GxP-M", "GxP-m2", "GxP-NA",
    "Kurzbeschreibung", "SW-Version / Typ:", "SW-Name:", "SW-Hersteller",
    "Bemerkung1", "Bemerkung2", "Bemerkung3", "Bemerkung4",
    "GxP_Produktqualitaet", "GxP_Patientensicherheit", "GxP_Datenintegritaet",
    "Systemtyp_CIS", "Systemtyp_CE", "Subtyp_PCS", "Subtyp_LCE", "Subtyp_EE",
    "VNAP_S0", "VNAP_S1", "VNAP_S2", "Subtyp_NA",
    "Offen", "Geschlossen", "NA",
    "KAT1", "KAT3", "KAT4", "KAT5", "KATNA",
    "Ersteller", "SME", "SI", "TSO", "BSO", "BQR", "CSQ",
    "Datum", "Version_Historie", "Historie", "Bearbeiter",
    "Hersteller", "Phenix", "SAP",
    "ERESTYP1", "ERESTYP2", "ERESTYP3", "ERESTYP4", "ERESTYPNA",
    "TTIEFEHOCH", "TTIEFEMITTEL", "TTIEFENIEDRIG",
    "Z1S1", "Z2S1", "Z3S1", "Z1S2", "Z1S3", "Z2S2", "Z2S3", "Z3S2", "Z3S3",
    "Lieferantennummer", "UeberlagerteMLCS", "Bedien-SOP", "SOP-Titel",
    "PLSTA", "DokNummerVorQualiPSO", "BE", "Raum", "Schnittstelle",
    "PNK", "BCkritisch", "BCunkritisch",
    "Neuerstellung", "Revisioniert",
    "GKATA", "GKATB1", "GKATB2", "GKATB3", "GKATB",
    "GKATC1", "GKATC2", "GKATC3", "GKATC", "GKATNA",
    "VQ", "NVQ", "QUAL", "VAL",
    "KI1", "KI2", "KI3", "KI4", "KI5", "KI6", "KINA",
    "Steuerung erfolgt über?",
    "Prozessbeschreibung", "Daten", "Parameter",
    "Alarme (GxP-relevant)", "Chargenprotokoll", "Audit Trail (AT)",
    "Benutzer-verwaltung?", "Schnittstellen mit PLS",
    "Angeschlossenes Equipment",
    "Hyperlink",
    "Sonstiges", "KI Bewertung", "Besonderheiten",
    "Erkannte_Version",
]

# ============================================================
# MAPPING: Checkboxen Kapitel 2
# V10: Testtiefe direkt in Zelle 4 (nicht in separater Tabelle!)
# ============================================================
# ============================================================
# VALIDIERUNG: Kategorien, bei denen genau EIN "r"-Wert erwartet wird
# ============================================================
# Jede Kategorie = (Anzeigename, [(Excel-Feldname, Anzeige-Label), ...])
# Ausgeschlossen (noch unvollständig getrackt, siehe help.txt):
#   - Testtiefe (Gering/Mittel/Hoch) - fehlende N/A-Spalte
#   - Validierung/Qualifizierung nach SOP (QUAL/VAL) - fehlende 3. SOP-Spalte
VALIDATION_KATEGORIEN_V8 = [
    ("Neuerstellung/Änderung", [
        ("Neuerstellung", "Neuerstellung"),
        ("Revisioniert", "Änderung – im Einsatz/Aktualisierung"),
    ]),
    ("Systemtyp (Zugangsbeschränkung)", [
        ("Offen", "offen"),
        ("Geschlossen", "geschlossen"),
        ("NA", "N/A"),
    ]),
    ("GxP-Relevanz", [
        ("GxP_Relevan_JA", "Ja"),
        ("GxP_Relevan_NEIN", "Nein"),
    ]),
    ("Business Kritisch", [
        ("BCkritisch", "Ja"),
        ("BCunkritisch", "Nein"),
    ]),
    ("GxP-Kritikalität", [
        ("GxP-C", "Critical"),
        ("GxP-M", "Major"),
        ("GxP-m2", "minor"),
        ("GxP-NA", "N/A"),
    ]),
    ("CS-Typ", [
        ("Systemtyp_CIS", "CIS"),
        ("Subtyp_LCE", "CE-LCE"),
        ("Subtyp_PCS", "CE-PCS"),
        ("Subtyp_EE", "CE-EE"),
        ("VNAP_S0", "S0"),
        ("VNAP_S1", "S1"),
        ("VNAP_S2", "S2"),
        ("Subtyp_NA", "N/A"),
    ]),
    ("ERES-Typ", [
        ("ERESTYP1", "Typ 1"),
        ("ERESTYP2", "Typ 2"),
        ("ERESTYP3", "Typ 3"),
        ("ERESTYP4", "Typ 4"),
        ("ERESTYPNA", "N/A"),
    ]),
    ("GAMP5 Software-Kategorie", [
        ("KAT1", "SW-Kat 1"),
        ("KAT3", "SW-Kat 3"),
        ("KAT4", "SW-Kat 4"),
        ("KAT5", "SW-Kat 5"),
        ("KATNA", "N/A"),
    ]),
    ("Gerätekategorie", [
        ("GKATA", "A"),
        ("GKATB", "B"),
        ("GKATC", "C"),
        ("GKATNA", "N/A"),
    ]),
    ("Vereinfachte Qualifizierung", [
        ("VQ", "Ja"),
        ("NVQ", "Nein"),
    ]),
    # V8 hat kein KI-Kapitel -> keine KI-Reifegrad-Kategorie
]

def validiere_kategorien(data, kategorien):
    """
    Prüft je Kategorie, ob genau EIN "r"-Wert (angekreuzt) vorhanden ist.
    c-Werte werden nicht angezeigt - nur Kategoriename, ausgewählter
    Wert (falls genau einer) und Status.
    """
    print("\n" + "="*55)
    print("✅ VALIDIERUNG: Vollständigkeitsprüfung")
    print("="*55)
    for name, optionen in kategorien:
        gefunden = [label for feld, label in optionen if data.get(feld) == "r"]
        if len(gefunden) == 1:
            print(f"  ✅ {name:<32} = \"{gefunden[0]}\"")
        elif len(gefunden) == 0:
            print(f"  ❗ {name:<32} = KEIN Wert ausgewählt (erwartet: genau 1)")
        else:
            werte = ", ".join(f'"{g}"' for g in gefunden)
            print(f"  ❌ {name:<32} = MEHRERE Werte: {werte} (erwartet: genau 1)")

CHECKBOX_MAPPING_V8 = {
    6: {
        6: {
            0: {0: "GxP-C", 1: "GxP-M", 2: "GxP-m2", 3: "GxP-NA"},
            1: {
                0: "Systemtyp_CIS", 1: "Subtyp_LCE", 2: "Subtyp_PCS",
                3: "Subtyp_EE",     4: "VNAP_S0",    5: "VNAP_S1",
                6: "VNAP_S2",       7: "Subtyp_NA",
            },
            2: {0: "ERESTYP1", 1: "ERESTYP2", 2: "ERESTYP3",
                3: "ERESTYP4", 4: "ERESTYPNA"},
            3: {0: "KAT1", 1: "KAT3", 2: "KAT4", 3: "KAT5", 4: "KATNA"},
            # V8: Testtiefe direkt in Zelle 4! (kein eigenes NA-Feld in der
            # Master-Excel für Testtiefe-N/A vorhanden -> Index 3 bewusst nicht gemappt,
            # "NA" ist für Kapitel-1 Systemtyp offen/geschlossen/N-A reserviert)
            4: {0: "TTIEFENIEDRIG", 1: "TTIEFEMITTEL", 2: "TTIEFEHOCH"},
            6: {0: "GKATA", 1: "GKATB", 2: "GKATC", 3: "GKATNA"},
            7: {0: "QUAL", 1: "VAL"},
            # V8: KEIN KI-Kapitel -> Zelle 8 (KI Reifegrad) existiert nicht, kein Mapping.
        },
        8: {
            0: {0: "BCkritisch", 1: "BCunkritisch"},
            6: {0: "VQ", 1: "NVQ"},
        },
    },
}

# ============================================================
# MAPPING: Textfelder → Excel-Spalten
# ============================================================
TEXT_FIELD_MAPPING_V8 = {
    "Bezeichnung des Equipments / Systemname": "AS/BDIS-Name",
    "Bezeichnung des Equipment / Systemname":  "AS/BDIS-Name",
    "Bezeichnung des Equipment/ Systemname":   "AS/BDIS-Name",
    "Bezeichnung des Equipments/ Systemname":  "AS/BDIS-Name",
    "Kurzbeschreibung / Verwendungszweck":     "Kurzbeschreibung",
    "Kurzbeschreibung/Verwendungszweck":       "Kurzbeschreibung",
    "Hersteller / SW-Ersteller / Lieferant":   "_hersteller_raw",
    "Hersteller/ SW-Ersteller/Lieferant":      "_hersteller_raw",
    "Einsatzbereich(e)/-ort(e)":               "_betrieb_raw",
    "Prozessbeschreibung":                     "Prozessbeschreibung",
    "Daten":                                   "Daten",
    "Parameter":                               "Parameter",
    "Alarme (GxP critical)":                   "Alarme (GxP-relevant)",
    "Alarme (GxP-relevant)":                   "Alarme (GxP-relevant)",
    "Chargenprotokoll":                        "Chargenprotokoll",
    "Audit Trail / Audit Trail Review":        "Audit Trail (AT)",
    "Audit Trail (AT)":                        "Audit Trail (AT)",
    "Audit-Trail / Audit-Trail-Review":        "Audit Trail (AT)",
    "Benutzerverwaltung":                      "Benutzer-verwaltung?",
    "Schnittstellen":                          "Schnittstellen mit PLS",
    "Angeschlossene Equipments":               "Angeschlossenes Equipment",
    "Angeschlossenes Equipment":               "Angeschlossenes Equipment",
    "Steuerung erfolgt über":                  "Steuerung erfolgt über?",
    "Bedien-SOP":                              "Bedien-SOP",
    "SOP-Titel":                               "SOP-Titel",
    "Sonstiges":                               "Sonstiges",
    "KI Bewertung":                            "KI Bewertung",
    "KI-Bewertung":                            "KI Bewertung",
}

# ============================================================
# HILFSFUNKTIONEN
# ============================================================
def get_checkbox_state_from_sdt(sdt):
    ns = {"w14": "http://schemas.microsoft.com/office/word/2010/wordml"}
    checkbox = sdt.find(".//w14:checkbox", ns)
    if checkbox is None:
        return None
    checked = checkbox.find(".//w14:checked", ns)
    if checked is None:
        return "c"
    val = checked.get("{http://schemas.microsoft.com/office/word/2010/wordml}val")
    return "r" if val == "1" else "c"

def parse_text_checkboxes(text):
    """
    Parst Text-Checkboxen im Serienbrief-Format, z.B.:
    'c Critical* c Major r minor c N/A' -> geordnete Liste von Zuständen.
    r = angekreuzt, c = leer. Reihenfolge entspricht der Reihenfolge im Text,
    genau wie bei echten SDT-Checkboxen.
    """
    if not text:
        return []
    marker_positions = list(re.finditer(r"(?:^|\s)([rc])(?=\s)", text))
    if not marker_positions:
        return []
    result = []
    for i, m in enumerate(marker_positions):
        marker = m.group(1)
        start = m.end()
        end = marker_positions[i + 1].start() if i + 1 < len(marker_positions) else len(text)
        label = text[start:end].strip()
        result.append({"state": marker, "label": label})
    return result

def get_checkboxes_from_cell(cell):
    result = []
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for sdt in cell._element.findall(".//w:sdt", ns):
        state = get_checkbox_state_from_sdt(sdt)
        if state is not None:
            result.append({"state": state})
    if result:
        return result
    # Fallback: keine echten Formularfelder gefunden -> Text-Checkboxen (Serienbrief) parsen
    text = get_cell_text(cell)
    return parse_text_checkboxes(text)

def get_cell_text(cell):
    parts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return " ".join(parts)

def text_matches(search_term, cell_text):
    s = search_term.lower().strip()
    c = cell_text.lower().strip()
    if s in c or c in s:
        return True
    s_clean = re.sub(r"[^a-z0-9äöüß]", "", s)
    c_clean = re.sub(r"[^a-z0-9äöüß]", "", c)
    return bool(s_clean and c_clean and
                (s_clean in c_clean or c_clean in s_clean))

def _is_boilerplate(text):
    patterns = [
        r"^weiter mit", r"^ende der systembewertung",
        r"^qualifizierung nach", r"^validierung nach",
        r"^qu-sop-", r"^qu-ope-", r"^qu-mt-",
        r"^fra-qu-", r"^keine mlcs", r"^\(info:",
        r"^begründung der gxp",
        r"^1\.systemname, mlcs",
        r"^anlagen-ids/equipment-nr\./qc-id$",
        r"^system-identifier",
    ]
    t = text.lower().strip()
    return any(re.match(p, t) for p in patterns)

# ============================================================
# ABLEITUNGS-FUNKTIONEN
# ============================================================
def parse_betrieb(raw):
    result = {}
    if not raw:
        return result
    m = re.search(r'\b([A-Z]\d+)\b', raw)
    if m:
        result["Gebaeude"] = m.group(1)
    if "/" in raw:
        betrieb = raw.split("/")[0].strip()
        result["Betrieb"] = betrieb
    elif "," in raw:
        parts = [p.strip() for p in raw.split(",")]
        nicht_gebaeude = [p for p in parts
                          if not re.match(r'^[A-Z]\d+$', p)]
        result["Betrieb"] = " ".join(nicht_gebaeude).strip()
    else:
        result["Betrieb"] = raw.strip()
    betrieb = result.get("Betrieb", "")
    if betrieb:
        result["API"] = betrieb.split()[0].strip().rstrip(",")
    return result

def parse_hersteller(raw):
    result = {}
    if not raw:
        return result

    # SW-Hersteller: OOZ-Nummern und Zeilenumbrüche entfernen
    # SW-Hersteller: Prefix + OOZ + QTP-Teile entfernen
    sw_hersteller = re.sub(r'\s*\n\s*OOZ[A-Z0-9]*', '', raw).strip()
    sw_hersteller = re.sub(r'\s*OOZ[A-Z0-9]+\s*', ' ', sw_hersteller).strip()
    sw_hersteller = re.sub(
        r'(PU-Lieferant|Systemsoftware|Lieferant)\s*:\s*',
        '', sw_hersteller, flags=re.IGNORECASE
    ).strip()
    sw_hersteller = re.sub(
        r'\s*/\s*QTP-Customer\s*ID\s*:.*$', '', sw_hersteller,
        flags=re.IGNORECASE
    ).strip()
    result["SW-Hersteller"] = sw_hersteller

    # Hersteller: Prefix wie "PU-Lieferant:", "Systemsoftware:" entfernen
    # und nur den Firmennamen behalten
    hersteller_raw = re.split(
        r'[/,]|QualiPSO|QTP|Phenix|Ursprünglich|OOZ', raw, maxsplit=1
    )[0]
    # Prefix wie "PU-Lieferant: " entfernen
    hersteller_raw = re.sub(
        r'^(PU-Lieferant|Systemsoftware|Lieferant|Hersteller)\s*:\s*',
        '', hersteller_raw, flags=re.IGNORECASE
    ).strip().rstrip(":")
    result["Hersteller"] = hersteller_raw.strip()

    # QualiPSO-ID oder QTP-Customer ID → Lieferantennummer
    m = re.search(
        r'(?:QualiPSO|QTP)[-\s]*(?:Customer\s*)?ID\s*:?\s*([A-Z0-9]+)',
        raw, re.IGNORECASE
    )
    if m:
        result["Lieferantennummer"] = m.group(1).strip()

    # Phenix-ID
    m = re.search(r'Phenix[-\s]*ID\s*:?\s*(?:ID\s*)?([0-9]+)',
                  raw, re.IGNORECASE)
    if m:
        result["Phenix"] = m.group(1).strip()
    else:
        # OOZ-Nummer direkt (ohne Prefix) → Phenix
        m = re.search(r'\b(OOZ[A-Z0-9]+)\b', raw, re.IGNORECASE)
        if m and "Lieferantennummer" not in result:
            result["Phenix"] = m.group(1).strip()

    return result

def parse_dok_nr(filename):
    m = re.search(r'(QU-[A-Z]+-\d+)', filename, re.IGNORECASE)
    return m.group(1) if m else None

def parse_sw_version(text):
    result = {}
    if not text:
        return result
    m = re.search(r'[Vv]ersion\s+([A-Z0-9][A-Z0-9._\-]+)', text)
    if m:
        result["SW-Version / Typ:"] = m.group(1)
    for pattern in [r'(800xA)', r'(Freelance)', r'(WinCC)',
                    r'(PCS\s*7)', r'(SCADA\s*\S+)', r'(TIA[-\s]*Portal)']:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            result["SW-Name:"] = m.group(1)
            break
    return result

def parse_gkat_subtypen(text):
    result = {}
    subtypen = {
        "GKATB1": r'B1', "GKATB2": r'B2', "GKATB3": r'B3',
        "GKATC1": r'C1', "GKATC2": r'C2', "GKATC3": r'C3',
    }
    for col, pattern in subtypen.items():
        result[col] = "r" if re.search(pattern, text or "") else "c"
    return result

def parse_gxp_einfluss(text):
    result = {
        "GxP_Produktqualitaet":    "",
        "GxP_Patientensicherheit": "",
        "GxP_Datenintegritaet":    "",
    }
    if not text or not text.strip():
        return result

    t = text.strip()
    hat_produkt = "produktqualität" in t.lower()
    hat_patient = "patientensicherheit" in t.lower()
    hat_daten   = ("datenintegrität" in t.lower() or
                   "datenintegrit" in t.lower())

    if hat_produkt or hat_patient or hat_daten:
        # Nur strukturiert splitten wenn Doppelpunkt vorhanden
        hat_doppelpunkt = bool(re.search(
            r'(?i)(produktqualität\s*:|patientensicherheit\s*:'
            r'|datenintegrit[äa]t\s*:)',
            t
        ))
        if not hat_doppelpunkt:
            result["GxP_Produktqualitaet"] = t
            return result

        teile = re.split(
            r'(?i)(produktqualität\s*[:\-]?|patientensicherheit\s*[:\-]?'
            r'|datenintegrit[äa]t\s*[:\-]?)',
            t
        )
        aktueller_key = None
        for teil in teile:
            tl = teil.lower().strip().rstrip(":").rstrip("-").strip()
            if tl == "produktqualität":
                aktueller_key = "GxP_Produktqualitaet"
            elif tl == "patientensicherheit":
                aktueller_key = "GxP_Patientensicherheit"
            elif tl in ("datenintegrität", "datenintegritaet"):
                aktueller_key = "GxP_Datenintegritaet"
            elif aktueller_key and teil.strip():
                # Komma-Artefakte am Anfang entfernen
                cleaned = teil.strip().lstrip(",").lstrip(".").strip()
                if cleaned:
                    result[aktueller_key] = (
                        result[aktueller_key] + " " + cleaned
                    ).strip()
    else:
        result["GxP_Produktqualitaet"] = t

    return result

def berechne_systemtyp_ce(data):
    if any(data.get(f) == "r"
           for f in ["Subtyp_PCS", "Subtyp_LCE", "Subtyp_EE"]):
        return "r"
    return "c"

# ============================================================
# EXTRAKTIONS-FUNKTIONEN
# ============================================================
def detect_template_version(doc):
    """
    Erkennt die Template-Version anhand struktureller Merkmale:
    V8:  Kein KI-Kapitel
    V10: KI-Kapitel + N/A bei Vereinfachte Qualifizierung
         ODER KI-Kapitel + Testtiefe-Eintrag in Kap. 2 (Checkbox oder Text)
    V11: GAMP 5 (2nd Edition) im Text
         ODER KI-Kapitel ohne obige V10-Merkmale
    """
    full_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += get_cell_text(cell).lower() + " "

    # Merkmal 1: GAMP 5 2nd Edition → eindeutig V11
    if "2nd edition" in full_text:
        return 11

    # Merkmal 2: Hat das Dokument ein KI-Kapitel?
    has_ki = any(k in full_text for k in [
        "künstliche intelligenz",
        "ki zum einsatz",
        "reifegrad",
        "autonomie",
        "steuerungsdesign",
    ])

    if not has_ki:
        return 8  # Kein KI-Kapitel → V8

    # Merkmal 3: Hat "Vereinfachte Qualifizierung" eine N/A-Option?
    has_vq_na = False
    for table in doc.tables:
        rows = list(table.rows)
        for row_idx, row in enumerate(rows):
            for cell in row.cells:
                text = get_cell_text(cell).lower()
                if "vereinfachte qualifizierung" in text:
                    for next_row in rows[row_idx+1:row_idx+6]:
                        for nc in next_row.cells:
                            nc_text = get_cell_text(nc).strip().lower()
                            cbs = get_checkboxes_from_cell(nc)
                            if nc_text in ["n/a", " n/a", "na"] and cbs:
                                has_vq_na = True

    if has_vq_na:
        return 10  # KI vorhanden + N/A → V10

    # Merkmal 4: Testtiefe in Kap. 2 vorhanden?
    # V10: Testtiefe steht in Zusammenfassungstabelle
    #      als echte Checkbox ODER als Text "r Gering/Mittel/Hoch/N/A"
    # V11: Testtiefe NUR als Matrix in Kap. 8
    has_testtiefe_in_kap2 = False
    for table in doc.tables:
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += get_cell_text(cell).lower() + " "

        # Suche Zusammenfassungstabelle (hat mlcs + testtiefe + gxp)
        if not all(k in table_text for k in [
                "mlcs", "testtiefe", "gxp"]):
            continue

        for row in table.rows:
            for cell in row.cells:
                cell_text = get_cell_text(cell).lower()
                # Echte Checkboxen in Testtiefe-Zelle
                if "testtiefe" in cell_text:
                    cbs = get_checkboxes_from_cell(cell)
                    if cbs:
                        has_testtiefe_in_kap2 = True
                # Text-Checkboxen: "r gering", "r mittel", "r hoch", "r n/a"
                if re.search(
                    r'\br\s+(gering|mittel|hoch|n/a)\b',
                    cell_text
                ):
                    has_testtiefe_in_kap2 = True

    if has_testtiefe_in_kap2:
        return 10  # Testtiefe in Kap. 2 → V10
    else:
        return 11  # Keine Testtiefe in Kap. 2 → V11
def extract_mlcs_id(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = get_cell_text(cells[0])
            if "system id" in label.lower() and "mlcs" in label.lower():
                seen = set()
                for i in range(len(cells)-1, 0, -1):
                    val = get_cell_text(cells[i]).strip()
                    if val in seen:
                        continue
                    seen.add(val)
                    val = val.replace("System ID (MLCS):", "").strip()
                    val = re.sub(r"(?i)^mlcs[\s\-:]*", "", val).strip()
                    val = re.sub(r"(?i)^id[\s\-:]*", "", val).strip()
                    val = re.sub(r"gemäß\s+\S+", "", val).strip()
                    val = re.sub(r"\(Info:.*?\)", "", val).strip()
                    val = re.sub(r"\(MLCS.*?\)", "", val,
                                 flags=re.IGNORECASE).strip()
                    if "system-identifier" in val.lower():
                        continue
                    if "cs inventarnummer" in val.lower():
                        continue
                    if val and not _is_boilerplate(val):
                        return val
    return None

def extract_ueberlagertes_mlcs(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = get_cell_text(cells[0])
            if ("schnittstelle zu anderen" in label.lower() and
                    "übergeordneten" in label.lower()):
                seen = set()
                for i in range(len(cells)-1, 0, -1):
                    val = get_cell_text(cells[i]).strip()
                    if val in seen:
                        continue
                    seen.add(val)
                    val = re.sub(
                        r"schnittstelle zu anderen/übergeordneten systemen:?",
                        "", val, flags=re.IGNORECASE
                    ).strip()
                    val = re.sub(r"(?i)^mlcs[\s\-:]*", "", val).strip()
                    val = re.sub(r"(?i)^id[\s\-:]*", "", val).strip()
                    if val and not _is_boilerplate(val):
                        return val
    return None

def extract_systemtyp_zugang(doc):
    """
    Liest Offen/Geschlossen/N-A aus der Systemtyp-Zeile in Kapitel 1
    (Zugangsbeschränkung nach 21 CFR Part 11, §11.3(4)).
    """
    result = {}
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = get_cell_text(cells[0])
            if "systemtyp" in label.lower() and "zugang" in label.lower():
                cbs = get_checkboxes_from_cell(cells[1])
                if len(cbs) >= 3:
                    result["Offen"]       = cbs[0]["state"]
                    result["Geschlossen"] = cbs[1]["state"]
                    result["NA"]          = cbs[2]["state"]
                return result
    return result

def extract_template_basis_version(doc):
    """
    Liest die Master-Template-Version, auf deren Basis dieses Dokument
    erstellt wurde, aus dem Text 'Neuerstellung auf Basis der ... Version
    X.Y' (Tabelle 'Version / Freigabedatum', Spalte 'Grund der Erstellung
    / Änderung'). Wichtig, weil manche Template-Versionen (z.B. V7 und V8)
    strukturell identisch sind - die reine Strukturerkennung würde beide
    als V8 einstufen, obwohl es sich um unterschiedliche, im Text explizit
    benannte Master-Template-Versionen handelt. Gibt z.B. "V7" zurück.
    """
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        header0 = get_cell_text(rows[0].cells[0]) if len(rows[0].cells) >= 1 else ""
        if "version" in header0.lower() and "freigabe" in header0.lower():
            if len(rows) > 1 and len(rows[1].cells) > 1:
                grund_text = get_cell_text(rows[1].cells[1])
                m = re.search(r"auf\s+basis\s+der.*?version\s+(\d+(?:\.\d+)?)",
                               grund_text, re.IGNORECASE | re.DOTALL)
                if m:
                    return f"V{m.group(1).split('.')[0]}"
    return None

def extract_version_freigabedatum(doc):
    """
    Liest die höchste Version aus der 'Version / Freigabedatum'-Tabelle auf
    der ersten Seite. Diese Tabelle kann mehrere Revisionszeilen enthalten;
    die aktuellste (höchste) Version steht nicht immer in der ersten
    Datenzeile, sondern manchmal weiter unten (z.B. Zeile 3).
    """
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue
        header0 = get_cell_text(rows[0].cells[0]) if len(rows[0].cells) >= 1 else ""
        if "version" in header0.lower() and "freigabe" in header0.lower():
            versionen = []
            for row in rows[1:]:
                if len(row.cells) < 1:
                    continue
                val = get_cell_text(row.cells[0]).strip()
                m = re.match(r"^(\d+\.\d+)", val)
                if m:
                    versionen.append(m.group(1))
            if versionen:
                return max(versionen, key=lambda v: [int(x) for x in v.split(".")])
    return None

def extract_anlage(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = get_cell_text(cells[0])
            if ("anlagen-id" in label.lower() or
                    "equipment-nr" in label.lower()):
                seen = set()
                for i in range(len(cells)-1, 0, -1):
                    val = get_cell_text(cells[i]).strip()
                    if val in seen:
                        continue
                    seen.add(val)
                    val = re.sub(
                        r"anlagen-ids/equipment-nr\./qc-id:?",
                        "", val, flags=re.IGNORECASE
                    ).strip()
                    if val and not _is_boilerplate(val):
                        return val
    return None

def extract_schnittstelle(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label = get_cell_text(cells[0])
            if "schnittstelle zu anderen" in label.lower():
                seen = set()
                for i in range(len(cells)-1, 0, -1):
                    val = get_cell_text(cells[i]).strip()
                    if val in seen:
                        continue
                    seen.add(val)
                    val = re.sub(
                        r"schnittstelle zu anderen.*?systemen:?",
                        "", val, flags=re.IGNORECASE
                    ).strip()
                    if val and not _is_boilerplate(val):
                        return val
    return None

def extract_besonderheiten(doc):
    """Besonderheiten-Text aus Kapitel 2 - sucht in allen Zellen der Zeile."""
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            for i, cell in enumerate(cells):
                text = get_cell_text(cell)
                if "besonder" in text.lower() and len(text) > 20:
                    # Suche in ALLEN folgenden Zellen der Zeile
                    for j in range(i + 1, len(cells)):
                        val = get_cell_text(cells[j])
                        if val and not _is_boilerplate(val):
                            return val
                    # Fallback: Text steht im Label selbst nach ":"
                    if ":" in text:
                        val = text.split(":", 1)[1].strip()
                        if val and not _is_boilerplate(val):
                            return val
    return None

def extract_begruendung_gxp(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 1:
                continue
            label = get_cell_text(cells[0])
            if ("begründung" in label.lower() and
                    "gxp" in label.lower()):
                if ":" in label:
                    text = label.split(":", 1)[1].strip()
                    if text:
                        return text
                if len(cells) > 1:
                    val = get_cell_text(cells[1])
                    if val and not _is_boilerplate(val):
                        return val
                return label
    return None

def extract_neuerstellung(doc):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = get_cell_text(cell)
                if ("neuerstellung" in text.lower() and
                        "änderung" in text.lower()):
                    checkboxes = get_checkboxes_from_cell(cell)
                    if len(checkboxes) >= 2:
                        return {
                            "Neuerstellung": checkboxes[0]["state"],
                            "Revisioniert":  checkboxes[1]["state"],
                        }
    return {}

def extract_gxp_relevan_bc(doc):
    result = {}
    keywords_gxp = ["gxp relevanz", "gxp-relevanz"]
    keywords_bc  = ["business kritikalität", "geschäftskritisch",
                    "business kritisch"]

    for table in doc.tables:
        rows = list(table.rows)
        for row_idx, row in enumerate(rows):
            for cell in row.cells:
                text = get_cell_text(cell).lower()

                if ("GxP_Relevan_JA" not in result and
                        any(k in text for k in keywords_gxp)):
                    for next_row in rows[row_idx+1:row_idx+6]:
                        for nc in next_row.cells:
                            nc_text = get_cell_text(nc).strip().lower()
                            # Marker-Präfix (r/c) aus Text-Checkboxen entfernen, z.B. "c ja" -> "ja"
                            nc_text_clean = re.sub(r"^[rc]\s+", "", nc_text).strip()
                            cbs = get_checkboxes_from_cell(nc)
                            if not cbs:
                                continue
                            if nc_text_clean == "ja":
                                result["GxP_Relevan_JA"]   = cbs[0]["state"]
                            elif nc_text_clean == "nein":
                                result["GxP_Relevan_NEIN"] = cbs[0]["state"]

                if ("BCkritisch" not in result and
                        any(k in text for k in keywords_bc)):
                    for next_row in rows[row_idx+1:row_idx+6]:
                        for nc in next_row.cells:
                            nc_text = get_cell_text(nc).strip().lower()
                            nc_text_clean = re.sub(r"^[rc]\s+", "", nc_text).strip()
                            cbs = get_checkboxes_from_cell(nc)
                            if not cbs:
                                continue
                            if nc_text_clean == "ja":
                                result["BCkritisch"]   = cbs[0]["state"]
                            elif nc_text_clean == "nein":
                                result["BCunkritisch"] = cbs[0]["state"]

    return result

def extract_text_fields(doc):
    result = {}
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            label_text = get_cell_text(cells[0])
            for search_term, excel_col in TEXT_FIELD_MAPPING_V8.items():
                if text_matches(search_term, label_text):
                    value = ""
                    for i in range(1, len(cells)):
                        candidate = get_cell_text(cells[i])
                        if candidate and not _is_boilerplate(candidate):
                            value = candidate
                            break
                    if value and excel_col not in result:
                        result[excel_col] = value
                    break
    return result

def extract_checkboxes_formularfelder(doc):
    result = {}
    for table_idx, table in enumerate(doc.tables):
        if table_idx not in CHECKBOX_MAPPING_V8:
            continue
        table_mapping = CHECKBOX_MAPPING_V8[table_idx]
        for row_idx, row in enumerate(table.rows):
            if row_idx not in table_mapping:
                continue
            row_mapping = table_mapping[row_idx]
            for cell_idx, cell in enumerate(row.cells):
                if cell_idx not in row_mapping:
                    continue
                cell_mapping = row_mapping[cell_idx]
                checkboxes   = get_checkboxes_from_cell(cell)
                for cb_idx, cb in enumerate(checkboxes):
                    if cb_idx in cell_mapping:
                        col = cell_mapping[cb_idx]
                        if col not in result:
                            result[col] = cb["state"]
    return result

def extract_history(doc):
    result = {}
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            col0 = get_cell_text(cells[0])
            col1 = get_cell_text(cells[1])
            if ("version" in col0.lower() and
                    "freigabe" in col0.lower()):
                continue
            m = re.match(r"^(\d+\.\d+)\s*/\s*(.+)$", col0.strip())
            if m:
                result["Version_Historie"] = m.group(1)
                result["Historie"]         = col1.strip()
                datum = m.group(2).strip()
                if "letztes unterschriftsdatum" not in datum.lower():
                    result["Datum"] = datum
    return result

# ============================================================
# EXCEL-OUTPUT
# ============================================================
# ============================================================
# MASTER-EXCEL: Konfiguration
# ============================================================
MASTER_EXCEL_PFAD   = r"C:\Users\de020409\Sanofi\FBC Betriebsübergreifende Dokumente - General\!Systembewertungen_CS\00_Serienbrief\Systembewertungen_GESAMT.xlsx"
MASTER_SHEET_NAME    = "SysBew"
MASTER_TABELLE_NAME  = "Tabelle1"

# Spalten, deren Name im Skript anders lautet als in der Master-Excel.
# Schlüssel = Feldname im Skript (data-Dict), Wert = Spaltenname in der Master-Excel.
MASTER_SPALTEN_MAPPING = {
    "Erkannte_Version": "Erkannte Version2",
}

def write_to_master_excel(data, docx_path, _versuch=1, _max_versuche=3):
    """
    Fügt die extrahierten Daten als neue Zeile in die Excel-Tabelle
    ('Tabelle1') der Master-Excel-Datei ein - per COM-Automatisierung
    (echtes Excel), NICHT per openpyxl.

    Grund: Die Master-Datei enthält eine Excel-Tabelle mit festem Bereich,
    ein Sensitivity-Label (Microsoft Purview Information Protection), eine
    externe Verknüpfung sowie Kommentare. openpyxl erhält beim Neuschreiben
    (load_workbook + save) weder die Tabellen-Bereichserweiterung noch das
    Sensitivity-Label noch die externe Verknüpfung zuverlässig -> Excel
    meldet beim nächsten Öffnen "Arbeitsmappe repariert". Über
    ListObject.ListRows.Add() erweitert echtes Excel die Tabelle korrekt und
    lässt alle anderen Dateibestandteile unangetastet.

    Bei transienten COM-Fehlern (z.B. RPC_E_DISCONNECTED, wenn Excel
    während des Vorgangs die Verbindung verliert - etwa durch OneDrive-
    Sync-Konflikte) wird bis zu _max_versuche-mal automatisch erneut
    versucht, ab dem zweiten Versuch immer mit einer komplett neuen,
    unsichtbaren Excel-Instanz (kein Anhängen an eine ggf. instabile
    laufende Instanz mehr).

    Benötigt: pip install pywin32   (nur unter Windows, mit installiertem Excel)
    """
    try:
        import win32com.client as win32
    except ImportError:
        print("\n  ❌ pywin32 ist nicht installiert.")
        print("     Bitte ausführen: pip install pywin32")
        return None

    if not os.path.exists(MASTER_EXCEL_PFAD):
        print(f"\n  ❌ Master-Excel nicht gefunden:")
        print(f"     {MASTER_EXCEL_PFAD}")
        print("     Bitte prüfen, ob OneDrive/SharePoint-Sync aktiv ist.")
        return None

    excel = None
    wb_com = None
    wir_haben_geoeffnet = False
    wir_haben_excel_gestartet = False

    try:
        try:
            if _versuch > 1:
                # Ab dem zweiten Versuch NIE an eine laufende Instanz anhängen -
                # die könnte genau die Ursache des vorherigen Fehlers sein.
                raise RuntimeError("Erzwinge neue Excel-Instanz für diesen Versuch")
            excel = win32.GetActiveObject("Excel.Application")
            _ = excel.Workbooks.Count  # Gesundheitscheck: haengt eine alte/verwaiste Instanz fest?
        except Exception:
            excel = win32.DispatchEx("Excel.Application")
            wir_haben_excel_gestartet = True

        excel.Visible = False
        excel.DisplayAlerts = False
        excel.AskToUpdateLinks = False

        ziel_pfad_abs = os.path.abspath(MASTER_EXCEL_PFAD).lower()
        for offene_wb in excel.Workbooks:
            try:
                if offene_wb.FullName.lower() == ziel_pfad_abs:
                    wb_com = offene_wb
                    break
            except Exception:
                continue

        if wb_com is None:
            wb_com = excel.Workbooks.Open(MASTER_EXCEL_PFAD, UpdateLinks=0, ReadOnly=False)
            wir_haben_geoeffnet = True

        ws_com = wb_com.Worksheets(MASTER_SHEET_NAME)

        try:
            tabelle = ws_com.ListObjects(MASTER_TABELLE_NAME)
        except Exception:
            print(f"\n  ❌ Excel-Tabelle '{MASTER_TABELLE_NAME}' nicht auf Sheet '{MASTER_SHEET_NAME}' gefunden.")
            if wir_haben_geoeffnet:
                wb_com.Close(SaveChanges=False)
            if wir_haben_excel_gestartet:
                excel.Quit()
            return None

        header_range = tabelle.HeaderRowRange
        header = [header_range.Cells(1, c).Value for c in range(1, header_range.Columns.Count + 1)]

        neue_zeile = tabelle.ListRows.Add()

        for col_idx, col_name in enumerate(header, start=1):
            if not col_name:
                continue
            script_key = col_name
            for skript_name, master_name in MASTER_SPALTEN_MAPPING.items():
                if master_name == col_name:
                    script_key = skript_name
                    break
            wert = data.get(script_key, "")
            neue_zeile.Range.Cells(1, col_idx).Value = wert

        ziel_zeile = neue_zeile.Range.Row
        wb_com.Save()

        # Duplikat-Hinweis (Dok.-Nr. + Version) für den soeben eingefügten Eintrag.
        # Rein informativ - darf den Haupt-Workflow niemals stören.
        try:
            dok_nr_col_idx  = header.index("Dok. -Nr.") + 1 if "Dok. -Nr." in header else None
            version_col_idx = header.index("Version") + 1 if "Version" in header else None
            if dok_nr_col_idx and version_col_idx:
                neuer_dok_nr = str(neue_zeile.Range.Cells(1, dok_nr_col_idx).Value or "").strip()
                neue_version = str(neue_zeile.Range.Cells(1, version_col_idx).Value or "").strip()
                PLATZHALTER = {"", "qu-ope-xxxxx", "n/a"}
                if neuer_dok_nr and neuer_dok_nr.lower() not in PLATZHALTER:
                    data_range = tabelle.DataBodyRange
                    dok_spalte = data_range.Columns(dok_nr_col_idx).Value
                    ver_spalte = data_range.Columns(version_col_idx).Value
                    if not isinstance(dok_spalte, tuple):
                        dok_spalte = ((dok_spalte,),)
                        ver_spalte = ((ver_spalte,),)
                    treffer = 0
                    for (dv,), (vv,) in zip(dok_spalte, ver_spalte):
                        if str(dv or "").strip() == neuer_dok_nr and str(vv or "").strip() == neue_version:
                            treffer += 1
                    if treffer > 1:
                        print(f"\n  ⚠️  HINWEIS: Möglicher Duplikat-Eintrag!")
                        print(f"     Dok.-Nr. '{neuer_dok_nr}' mit Version '{neue_version}' kommt")
                        print(f"     bereits {treffer}x in der Master-Excel vor (inkl. diesem neuen Eintrag).")
        except Exception:
            pass

        if wir_haben_geoeffnet:
            wb_com.Close(SaveChanges=False)  # bereits gespeichert
        if wir_haben_excel_gestartet:
            excel.Quit()

        if _versuch > 1:
            print(f"     ✅ Erfolgreich im {_versuch}. Versuch.")

        return ziel_zeile

    except Exception as e:
        print(f"\n  ❌ Fehler beim Schreiben in die Master-Excel (COM): {e}")
        try:
            if wir_haben_geoeffnet and wb_com is not None:
                wb_com.Close(SaveChanges=False)
            if wir_haben_excel_gestartet and excel is not None:
                excel.Quit()
        except Exception:
            pass

        if _versuch < _max_versuche:
            print(f"     Erneuter Versuch ({_versuch + 1}/{_max_versuche}) mit frischer Excel-Instanz...")
            time.sleep(3)
            return write_to_master_excel(data, docx_path, _versuch=_versuch + 1, _max_versuche=_max_versuche)

        print("     Falls dies ein 'OLE error 0x800a01a8' o.ä. ist: im Task-")
        print("     Manager prüfen, ob ein unsichtbarer EXCEL.EXE-Prozess von")
        print("     einem früheren Lauf noch offen ist, diesen beenden und")
        print("     erneut versuchen.")
        return None

# ============================================================
# HAUPTFUNKTION
# ============================================================
def parse_systembewertung_v8(docx_path):
    # Pfad bereinigen (Anführungszeichen vom Drag & Drop entfernen)
    docx_path = docx_path.strip().strip('"').strip("'")

    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Datei nicht gefunden: {docx_path}")

    ext = os.path.splitext(docx_path)[1].lower()
    if ext not in [".docx", ".doc"]:
        print(f"\n  ❌ Keine Word-Datei: '{os.path.basename(docx_path)}'")
        print(f"     Nur .docx Dateien werden unterstützt.")
        return {}

    # Prüfe ob Datei lokal verfügbar (kein OneDrive-Platzhalter)
    try:
        with open(docx_path, 'rb') as f:
            header = f.read(4)
        if header[:2] != b'PK':
            print(f"\n  ❌ Datei ist kein gültiges .docx!")
            print(f"     Möglicherweise nur OneDrive-Platzhalter.")
            print(f"     → Rechtsklick → 'Immer auf diesem Gerät behalten'")
            return {}
    except PermissionError:
        print(f"\n  ❌ Datei ist gesperrt!")
        print(f"     ⚠️  WICHTIG: Bitte Word KOMPLETT schließen!")
        print(f"     → Alle Word-Fenster schließen")
        print(f"     → Task-Manager prüfen (Ctrl+Shift+Esc)")
        print(f"     → Nach WINWORD.EXE suchen und beenden")
        print(f"     → Dann Skript erneut starten")
        return {}
    except OSError as e:
        print(f"\n  ❌ Datei-Zugriffsfehler: {e}")
        print(f"     OneDrive-Synchronisation abwarten und erneut versuchen.")
        return {}

    print(f"\n📄 Lese Datei: {os.path.basename(docx_path)}")
    doc  = Document(docx_path)
    data = {}

    print("  → Prüfe Template-Version...")
    template_version = detect_template_version(doc)

    if template_version == 8:
        print("  ✅ Template-Version: V8 — passt zu diesem Script")
        data["Erkannte_Version"] = "V8"
    else:
        if template_version is not None:
            print(f"  ❌ Template-Version: V{template_version}")
            print(f"     Dieses Script ist nur für V8!")
            print(f"     Bitte word_parser_v{template_version}.py verwenden.")
        else:
            print("  ❌ Template-Version: nicht erkannt")
            print("     Dieses Script ist nur für V8!")
        print("\n  ⛔ Verarbeitung abgebrochen.")
        return {}

    print("  → Extrahiere MLCS-ID...")
    mlcs_id = extract_mlcs_id(doc)
    if mlcs_id:
        data["MLCSID"] = mlcs_id

    print("  → Extrahiere Anlage/Equipment-Nr...")
    anlage = extract_anlage(doc)
    if anlage:
        data["Anlage"] = anlage

    print("  → Extrahiere Schnittstelle...")
    schnittstelle = extract_schnittstelle(doc)
    if schnittstelle:
        data["Schnittstelle"] = schnittstelle

    print("  → Extrahiere überlagertes MLCS...")
    ueberlagertes_mlcs = extract_ueberlagertes_mlcs(doc)
    if ueberlagertes_mlcs:
        data["UeberlagerteMLCS"] = ueberlagertes_mlcs

    print("  → Extrahiere Offen/Geschlossen/N-A (Kapitel 1 Systemtyp)...")
    data.update(extract_systemtyp_zugang(doc))

    print("  → Extrahiere Textfelder...")
    data.update(extract_text_fields(doc))

    print("  → Verarbeite Betrieb / API / Gebäude...")
    betrieb_raw = data.pop("_betrieb_raw", None)
    if betrieb_raw:
        data.update(parse_betrieb(betrieb_raw))

    print("  → Verarbeite Hersteller / Lieferant...")
    hersteller_raw = data.pop("_hersteller_raw", None)
    if hersteller_raw:
        data.update(parse_hersteller(hersteller_raw))

    print("  → Extrahiere Dok.-Nr. aus Dateiname...")
    dok_nr = parse_dok_nr(os.path.basename(docx_path))
    if dok_nr:
        data["Dok. -Nr."] = dok_nr

    print("  → Extrahiere Checkboxen (Kapitel 2 inkl. Testtiefe)...")
    data.update(extract_checkboxes_formularfelder(doc))

    print("  → Extrahiere Neuerstellung/Revisioniert (text-basiert)...")
    data.update(extract_neuerstellung(doc))

    print("  → Extrahiere GxP-Relevanz + BC (text-basiert)...")
    data.update(extract_gxp_relevan_bc(doc))

    print("  → Berechne Systemtyp_CE...")
    data["Systemtyp_CE"] = berechne_systemtyp_ce(data)

    print("  → Extrahiere Besonderheiten / SW-Version / GKAT...")
    besonderheiten = extract_besonderheiten(doc)
    if besonderheiten:
        data["Besonderheiten"] = besonderheiten
        data.update(parse_sw_version(besonderheiten))
        data.update(parse_gkat_subtypen(besonderheiten))

    print("  → Analysiere GxP-Begründung...")
    begruendung = extract_begruendung_gxp(doc)
    data.update(parse_gxp_einfluss(begruendung))

    print("  → Extrahiere Dokumentenhistorie...")
    data.update(extract_history(doc))

    version_freigabe = extract_version_freigabedatum(doc)
    data["Version"] = version_freigabe if version_freigabe else data.get("Version_Historie", "")

    # Präzisere Erkannte_Version: manche Master-Templates (z.B. V7) sind
    # strukturell identisch zu V8 - der Basis-Versionshinweis im Text ist
    # genauer als die reine Strukturerkennung.
    basis_version = extract_template_basis_version(doc)
    if basis_version:
        data["Erkannte_Version"] = basis_version

    print("  → Setze Bemerkungen...")
    data["Bemerkung1"] = data.get("Prozessbeschreibung", "")
    data["Bemerkung2"] = data.get("Daten", "")
    data["Bemerkung3"] = data.get("Audit Trail (AT)", "")
    data["Bemerkung4"] = data.get("Parameter", "")

    # V8 hat kein KI-Kapitel -> Feld fest mit Hinweistext befüllen (überschreibt
    # bewusst jeden ggf. zufällig gefundenen Wert, da V8-Dokumente kein KI-Kapitel haben)
    data["KI Bewertung"] = "keine KI-Bewertung vorhanden"

    print(f"  ✅ Extraktion abgeschlossen: {len(data)} Felder gefunden")
    return data

# ============================================================
# START: Drag & Drop ODER Kommandozeile
# ============================================================
if __name__ == "__main__":

    if len(sys.argv) >= 2:
        docx_path = sys.argv[1]
    else:
        print("="*55)
        print("  V8 Systembewertung Extraktor")
        print("="*55)
        print("  ⚠️  WICHTIG: Word-Datei muss geschlossen sein!")
        print("="*55)
        print("  Tipp: Word-Datei auf dieses Script ziehen!")
        print("="*55)
        docx_path = input("\n📂 Pfad zur Word-Datei: ").strip().strip('"')
    try:
        data = parse_systembewertung_v8(docx_path)

        if not data:
            print("\n  Keine Daten extrahiert — kein Excel erstellt.")
        else:
            print("\n" + "="*55)
            print("📊 EXTRAHIERTE DATEN (Vorschau):")
            print("="*55)
            for col, value in sorted(data.items()):
                display_val = (
                    str(value)[:60] + "..."
                    if len(str(value)) > 60
                    else str(value)
                )
                print(f"  {col:<35} = {display_val}")

            validiere_kategorien(data, VALIDATION_KATEGORIEN_V8)

            print("\n  → Füge Zeile in Master-Excel ein...")
            ziel_zeile = write_to_master_excel(data, docx_path)
            if ziel_zeile:
                print(f"\n✅ In Master-Excel eingefügt (Zeile {ziel_zeile}):")
                print(f"   {MASTER_EXCEL_PFAD}")
            else:
                print("\n❌ Master-Excel konnte nicht aktualisiert werden.")

    except FileNotFoundError as e:
        print(f"\n❌ Fehler: {e}")
    except Exception as e:
        print(f"\n❌ Unerwarteter Fehler: {e}")
        import traceback
        traceback.print_exc()

    print("\n" + "="*55)
    print("  Drücke ENTER zum Beenden...")
    input()
